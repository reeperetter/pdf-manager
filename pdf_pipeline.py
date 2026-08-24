"""
Основний PDF-пайплайн: рендеринг сторінок, вставка розпізнаного/
перекладеного тексту в PDF, конвертація зображень у PDF, збереження
окремого текстового файлу (DOCX/TXT).

Сам рушій розпізнавання (Tesseract) та розпрямлення сторінок - у ocr.py.
"""
import gc
import io
import os
import random
import time

import docx as _docx_module
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from PIL import Image, ImageOps

from ocr import (detect_and_fix_orientation, dewarp_page_image,
                 ocr_lines_from_image)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BUNDLED_FONT = os.path.join(SCRIPT_DIR, "DejaVuSans.ttf")

FALLBACK_FONTS = [
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",         # шлях на Fedora/Nobara
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/System/Library/Fonts/Helvetica.ttc",
]


class ProcessingCancelled(Exception):
    pass


def interruptible_sleep(duration, cancel_event=None, step=0.05):
    """Виконує паузу з можливістю миттєвого переривання через cancel_event."""
    end_time = time.time() + duration
    while time.time() < end_time:
        if cancel_event and cancel_event.is_set():
            raise ProcessingCancelled()
        time.sleep(min(step, max(0, end_time - time.time())))


def find_unicode_font(user_path=None):
    """Повертає шлях до TTF-шрифту, що підтримує кирилицю."""
    if user_path and os.path.isfile(user_path):
        return user_path
    if os.path.isfile(BUNDLED_FONT):
        return BUNDLED_FONT
    for p in FALLBACK_FONTS:
        if os.path.isfile(p):
            return p
    return None


def extract_lines_from_page(page):
    """Витягує рядки тексту з наявного текстового шару PDF."""
    lines = []
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            text = "".join(s["text"] for s in spans).strip()
            if not text:
                continue
            x0 = min(s["bbox"][0] for s in spans)
            y0 = min(s["bbox"][1] for s in spans)
            x1 = max(s["bbox"][2] for s in spans)
            y1 = max(s["bbox"][3] for s in spans)
            lines.append(
                {"text": text, "bbox": (x0, y0, x1, y1), "page_coords": True}
            )
    return lines


def line_to_rect(line, scale_x, scale_y):
    """Перетворює bbox рядка у fitz.Rect."""
    x0, y0, x1, y1 = line["bbox"]
    if line.get("page_coords"):
        return fitz.Rect(x0, y0, x1, y1)
    return fitz.Rect(x0 * scale_x, y0 * scale_y, x1 * scale_x, y1 * scale_y)


def line_to_pixel_bbox(line, scale_x, scale_y):
    """Перетворює bbox рядка у координати пікселів зображення."""
    x0, y0, x1, y1 = line["bbox"]
    if line.get("page_coords"):
        return (x0 / scale_x, y0 / scale_y, x1 / scale_x, y1 / scale_y)
    return line["bbox"]


def fit_fontsize(text, rect_width, rect_height, fontfile, max_size=None):
    """Підбирає розмір шрифту під прямокутник."""
    if max_size is None:
        max_size = max(4, rect_height * 0.85)
    size = max_size
    try:
        font = fitz.Font(fontfile=fontfile)
    except Exception:
        font = fitz.Font("helv")
    while size > 4:
        width = font.text_length(text, fontsize=size)
        if width <= rect_width or size <= 4.5:
            break
        size -= 0.5
    return size


def sample_bg_color(pil_img, bbox, pad=4):
    """Визначає колір фону навколо текстового блоку."""
    x0, y0, x1, y1 = bbox
    w, h = pil_img.size
    x0i, y0i, x1i, y1i = int(x0), int(y0), int(x1), int(y1)

    strips = [
        (x0i, max(0, y0i - pad), x1i, y0i),
        (x0i, y1i, x1i, min(h, y1i + pad)),
        (max(0, x0i - pad), y0i, x0i, y1i),
        (x1i, y0i, min(w, x1i + pad), y1i),
    ]

    samples_r, samples_g, samples_b = [], [], []
    for sx, sy, ex, ey in strips:
        if ex <= sx or ey <= sy:
            continue
        crop = pil_img.crop((sx, sy, ex, ey)).convert("RGB")
        for p in crop.getdata():
            samples_r.append(p[0])
            samples_g.append(p[1])
            samples_b.append(p[2])

    if not samples_r:
        return (1, 1, 1)

    def median(vals):
        vals = sorted(vals)
        n = len(vals)
        mid = n // 2
        if n % 2:
            return vals[mid]
        return (vals[mid - 1] + vals[mid]) / 2

    return (median(samples_r) / 255, median(samples_g) / 255, median(samples_b) / 255)


IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp")


def images_to_pdf(image_paths, output_path, log_fn=None, progress_fn=None,
                  cancel_event=None, jpeg_quality=88, max_dimension=2200):
    """Збирає список зображень в один PDF."""
    assumed_dpi = 200
    doc = fitz.open()
    total_before = 0
    total_after = 0
    try:
        for i, img_path in enumerate(image_paths):
            if cancel_event is not None and cancel_event.is_set():
                raise ProcessingCancelled()
            try:
                total_before += os.path.getsize(img_path)
            except OSError:
                pass

            img = Image.open(img_path)
            img = ImageOps.exif_transpose(img)
            img = img.convert("RGB")

            if max(img.width, img.height) > max_dimension:
                scale = max_dimension / max(img.width, img.height)
                img = img.resize(
                    (max(1, int(img.width * scale)),
                     max(1, int(img.height * scale))),
                    Image.LANCZOS,
                )

            w_pt = img.width * 72 / assumed_dpi
            h_pt = img.height * 72 / assumed_dpi
            page = doc.new_page(width=w_pt, height=h_pt)

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
            total_after += buf.tell()
            page.insert_image(fitz.Rect(0, 0, w_pt, h_pt),
                              stream=buf.getvalue())

            if log_fn:
                log_fn(
                    f"  Додано сторінку {i + 1}/{len(image_paths)}: "
                    f"{os.path.basename(img_path)}"
                )
            if progress_fn:
                progress_fn(i + 1, len(image_paths))
        doc.save(output_path, garbage=3, deflate=True)
    finally:
        doc.close()

    if log_fn and total_before > 0:
        final_size = os.path.getsize(output_path)
        log_fn(
            f"  Стиснення: вихідні зображення ~{total_before / 1_048_576:.1f} МБ "
            f"→ підсумковий PDF {final_size / 1_048_576:.1f} МБ"
        )
    return output_path


_ERROR_PAGE_MARKERS = (
    "that's an error",
    "that\u2019s an error",
    "that's all we know",
    "that\u2019s all we know",
    "please try again later",
    "<html",
    "error 500",
    "error 429",
    "500.that",
)


def _looks_like_error_page(text):
    if not text:
        return False
    low = text.lower()
    return any(marker in low for marker in _ERROR_PAGE_MARKERS)


def safe_translate(translator, text, log_fn=None, cancel_event=None, max_retries=4, base_delay=1.2):
    """Перекладає рядок з можливістю скасування та ретраями."""
    if not text or not text.strip():
        return text
    last_err = None
    for attempt in range(max_retries):
        if cancel_event and cancel_event.is_set():
            raise ProcessingCancelled()
        try:
            result = translator.translate(text)
        except Exception as e:
            last_err = e
            result = None
        if result and not _looks_like_error_page(result):
            return result
        if result and _looks_like_error_page(result):
            last_err = RuntimeError("Google повернув сторінку помилки замість перекладу")

        delay = base_delay * (2 ** attempt) + random.uniform(0, 0.5)
        interruptible_sleep(delay, cancel_event=cancel_event)

    if log_fn:
        log_fn(
            f"  [!] Не вдалося перекласти рядок після {max_retries} спроб "
            f"({last_err}); залишаю оригінальний текст."
        )
    return text


def _write_txt_file(path, page_texts):
    with open(path, "w", encoding="utf-8") as f:
        for i, text in enumerate(page_texts):
            if len(page_texts) > 1:
                if i > 0:
                    f.write("\n\n")
                f.write(f"----- Сторінка {i + 1} -----\n\n")
            f.write(text)
            f.write("\n")


def _write_docx_file(path, page_texts):
    doc = _docx_module.Document()
    for i, text in enumerate(page_texts):
        if i > 0:
            doc.add_page_break()
        if len(page_texts) > 1:
            heading = doc.add_paragraph(f"Сторінка {i + 1}")
            heading.runs[0].bold = True
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
    doc.save(path)


def save_text_outputs(page_texts, output_dir, base_name, suffix, formats, split_pages):
    saved = []
    writers = {"txt": _write_txt_file, "docx": _write_docx_file}
    if split_pages:
        for i, text in enumerate(page_texts):
            for fmt in formats:
                out_path = os.path.join(
                    output_dir, f"{base_name}{suffix}_p{i + 1}.{fmt}")
                writers[fmt](out_path, [text])
                saved.append(out_path)
    else:
        for fmt in formats:
            out_path = os.path.join(output_dir, f"{base_name}{suffix}.{fmt}")
            writers[fmt](out_path, page_texts)
            saved.append(out_path)
    return saved


def process_pdf(
    input_path,
    output_dir,
    dpi,
    make_searchable,
    translate_targets,
    font_path,
    log_fn,
    cancel_event,
    page_start=None,
    page_end=None,
    page_progress_fn=None,
    glossary=None,
    low_confidence_threshold=55,
    dewarp=False,
    export_text_formats=None,
    export_text_split_pages=False,
):
    glossary = glossary or set()
    low_confidence_pages = []
    export_text_formats = export_text_formats or set()
    original_page_texts = []
    translated_page_texts = {lang: [] for lang in translate_targets}

    base_name = os.path.splitext(os.path.basename(input_path))[0]
    doc = fitz.open(input_path)
    n_pages = doc.page_count

    fontfile = find_unicode_font(font_path)
    if fontfile is None:
        raise RuntimeError("Не знайдено TTF-шрифт з підтримкою кирилиці.")

    out_searchable = fitz.open() if make_searchable else None
    out_translated = {lang: fitz.open() for lang in translate_targets}

    translators = {
        lang: GoogleTranslator(source="auto", target=lang) for lang in translate_targets
    }
    translation_cache = {}

    for page_index in range(n_pages):
        if cancel_event and cancel_event.is_set():
            raise ProcessingCancelled()

        page_num = page_index + 1
        in_range = (page_start is None or page_num >= page_start) and \
                   (page_end is None or page_num <= page_end)

        log_fn(f"[{base_name}] Сторінка {page_num}/{n_pages}: рендеринг...")
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        pil_img = Image.open(io.BytesIO(img_bytes))

        if in_range:
            pil_img, was_rotated = detect_and_fix_orientation(
                pil_img, log_fn=log_fn)
            if was_rotated:
                buf = io.BytesIO()
                pil_img.save(buf, format="PNG")
                img_bytes = buf.getvalue()

        page_w = pil_img.width * 72 / dpi
        page_h = pil_img.height * 72 / dpi
        scale_x = page_w / pil_img.width
        scale_y = page_h / pil_img.height

        if dewarp and in_range:
            try:
                corrected = dewarp_page_image(
                    pil_img, log_fn=log_fn, cancel_event=cancel_event)
                if corrected is not pil_img:
                    pil_img = corrected
                    buf = io.BytesIO()
                    pil_img.save(buf, format="PNG")
                    img_bytes = buf.getvalue()
            except RuntimeError as e:
                log_fn(f"  [!] Розпрямлення недоступне: {e}")
                dewarp = False

        lines = []
        if in_range and (make_searchable or translate_targets):
            lines = extract_lines_from_page(page)
            if lines:
                log_fn(
                    f"[{base_name}] Сторінка {page_num}/{n_pages}: "
                    f"текст з PDF ({len(lines)} рядків), OCR пропущено")
            else:
                log_fn(
                    f"[{base_name}] Сторінка {page_num}/{n_pages}: "
                    f"немає текстового шару — розпізнавання (OCR)...")
                lines, ocr_score = ocr_lines_from_image(pil_img)
                if lines and ocr_score < low_confidence_threshold:
                    low_confidence_pages.append(page_num)
                    log_fn(
                        f"  [!] Сторінка {page_num}: невисока впевненість "
                        f"OCR ({ocr_score:.0f}/100) - варто перевірити вручну")
        elif not in_range:
            log_fn(
                f"[{base_name}] Сторінка {page_num}/{n_pages}: поза вибраним "
                f"діапазоном, пропущено (сторінку скопійовано без обробки)")

        if export_text_formats:
            original_page_texts.append("\n".join(ln["text"] for ln in lines))

        # 1) Searchable PDF
        if out_searchable is not None:
            sp = out_searchable.new_page(width=page_w, height=page_h)
            sp.insert_image(sp.rect, stream=img_bytes)
            for line in lines:
                rect = line_to_rect(line, scale_x, scale_y)
                if rect.is_empty or rect.width <= 0 or rect.height <= 0:
                    continue
                fs = fit_fontsize(
                    line["text"], rect.width, rect.height, fontfile)
                baseline = fitz.Point(rect.x0, rect.y1 - 0.15 * fs)
                try:
                    sp.insert_text(
                        baseline,
                        line["text"],
                        fontsize=fs,
                        fontfile=fontfile,
                        fontname="cyr1",
                        render_mode=3,
                    )
                except Exception as e:
                    log_fn(f"  [!] Пропущено рядок (текстовий шар): {e}")

        # 2) Translated PDF
        if translate_targets:
            translations = {}
            if lines:
                log_fn(
                    f"[{base_name}] Сторінка {page_num}/{n_pages}: переклад...")
                texts_to_translate = [ln["text"] for ln in lines]
                for lang in translate_targets:
                    result = []
                    for t in texts_to_translate:
                        if cancel_event and cancel_event.is_set():
                            raise ProcessingCancelled()

                        if t.strip().lower() in glossary:
                            result.append(t)
                            continue
                        cache_key = (lang, t)
                        if cache_key in translation_cache:
                            result.append(translation_cache[cache_key])
                            continue

                        tr = safe_translate(
                            translators[lang], t, log_fn, cancel_event=cancel_event)
                        translation_cache[cache_key] = tr
                        result.append(tr)

                        interruptible_sleep(0.2, cancel_event=cancel_event)
                    translations[lang] = result

            if export_text_formats:
                for lang in translate_targets:
                    text = "\n".join(translations.get(
                        lang, [])) if lines else ""
                    translated_page_texts[lang].append(text)

            for lang in translate_targets:
                tp = out_translated[lang].new_page(width=page_w, height=page_h)
                tp.insert_image(tp.rect, stream=img_bytes)
                if not lines:
                    continue
                translated_lines = translations[lang]

                to_draw = []
                for line, tr_text in zip(lines, translated_lines):
                    orig_text = line["text"].strip()
                    tr_text_stripped = (tr_text or "").strip()
                    if tr_text_stripped.lower() == orig_text.lower():
                        continue
                    rect = line_to_rect(line, scale_x, scale_y)
                    if rect.is_empty or rect.width <= 0 or rect.height <= 0:
                        continue
                    px_bbox = line_to_pixel_bbox(line, scale_x, scale_y)
                    bg = sample_bg_color(pil_img, px_bbox)
                    cover = fitz.Rect(rect.x0 - 1, rect.y0 - 1,
                                      rect.x1 + 1, rect.y1 + 1)
                    tp.draw_rect(cover, color=bg, fill=bg, overlay=True)
                    to_draw.append((rect, tr_text))

                for rect, tr_text in to_draw:
                    if not tr_text:
                        continue
                    fs = fit_fontsize(tr_text, rect.width,
                                      rect.height, fontfile)
                    baseline = fitz.Point(rect.x0, rect.y1 - 0.15 * fs)
                    try:
                        tp.insert_text(
                            baseline,
                            tr_text,
                            fontsize=fs,
                            fontfile=fontfile,
                            fontname="cyr1",
                            color=(0, 0, 0),
                        )
                    except Exception as e:
                        log_fn(f"  [!] Пропущено рядок (переклад {lang}): {e}")

        del pix, pil_img, img_bytes, lines
        page = None
        fitz.TOOLS.store_shrink(100)
        gc.collect()

        if page_progress_fn:
            page_progress_fn(page_num, n_pages)

    os.makedirs(output_dir, exist_ok=True)
    saved_files = []

    if out_searchable is not None:
        out_path = os.path.join(output_dir, f"{base_name}_searchable.pdf")
        out_searchable.save(out_path, garbage=3, deflate=True)
        out_searchable.close()
        saved_files.append(out_path)

    for lang in translate_targets:
        out_path = os.path.join(output_dir, f"{base_name}_{lang}.pdf")
        out_translated[lang].save(out_path, garbage=3, deflate=True)
        out_translated[lang].close()
        saved_files.append(out_path)

    if export_text_formats:
        if make_searchable:
            saved_files += save_text_outputs(
                original_page_texts, output_dir, base_name, "_text",
                export_text_formats, export_text_split_pages)
        for lang in translate_targets:
            saved_files += save_text_outputs(
                translated_page_texts[lang], output_dir, base_name, f"_{lang}_text",
                export_text_formats, export_text_split_pages)

    doc.close()
    gc.collect()
    return saved_files, low_confidence_pages